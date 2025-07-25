from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# Profile picture
document.add_picture(
    'Python.png',
    width=Inches(2.0)
)

# Name, phone number, and email details
name = input('What is your name? ')
speak('Hello ' + name + 'thank you for entering your name')
speak('Hello ' + name + 'Please input phone number? ')
phone_number = input('What is your phone number? ')
speak('Hello ' + name + 'thank you for entering your number ' + phone_number)
speak('Hello ' + name + 'please enter your email ')
email = input('What is your email? ')
speak('Hello ' + name + 'thank you for entering your email' + email)

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email  # Added spaces for readability
)

# About me
document.add_heading('About Me')
about_me = input('Tell me about yourself: ')
document.add_paragraph(about_me)

# Work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input('Enter Company: ')
from_date = input('From Date: ')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# More experiences
while True:
    has_more_experiences = input('Do you have more experience? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company: ')
        from_date = input('From Date: ')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break


 # SKILLS (Put this outside the while loop — after the experience loop ends)
document.add_heading('Skills')
skill = input('Enter a skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter another skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV made using help from Nelson at Amigoscode course project"

# Save the document
document.save('cv.docx')